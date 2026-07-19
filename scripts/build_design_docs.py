from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TROOT = Path('/Users/doublegreen/.codex/plugins/cache/openai-curated-remote/openai-templates/0.1.0/skills')
OUT = ROOT / 'artifacts'
OUT.mkdir(exist_ok=True)

def replace_nonempty(doc, values):
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    for p, value in zip(paragraphs, values):
        p.text = value

def design_report():
    ref = TROOT / 'artifact-template-design-report/assets/reference.docx'
    doc = Document(ref)
    vals = [
        'CHUTE Design Evolution Report', 'Contents', 'Executive summary',
        'The supplied reference is within reach as a disciplined translation: CHUTE can evolve from an arcade quiz into a compact match-intelligence terminal while preserving its Neo Arcade identity, quiz semantics, TxLINE evidence and Solana proof layer.',
        'At a glance', 'The highest-leverage move is a component-level redesign: dominant match metric, one decisive action, evidence-backed chart, accessible sheet primitives and restrained motion.',
        'Current scope', 'CHUTE already has consolidated tokens, replay and predictive modes, ranking, wallet flow and proof metadata. The missing layer is composition and interaction polish, not core product capability.',
        'Key findings', 'The reference uses oversized numeric hierarchy, fluorescent lime as a decision color, stacked cards, compact navigation and analytical charts. These patterns map naturally to CHUTE Matchday, Quiz, Reveal and Result screens.',
        'Implications', 'The current angular Neo Arcade language should remain for identity and competition. Rounded terminal surfaces can be introduced for data and interaction, producing a hybrid system rather than a generic fintech clone.',
        'Key takeaway. Proceed with a component-level redesign, not a wholesale framework replacement.',
        'Recommendations', 'Adopt Motion for gestures and transitions, Radix primitives for accessible behavior-heavy components, Recharts for evidence-backed charts, and consolidate the current icon overlap during migration.',
        'Phase A — surface foundation. Add terminal surfaces, MetricHero, ThesisCard and ProofCapsule.',
        'Phase B — analytical interaction. Add snapshot-backed SignalChart, selected-card transitions and accessible wallet/network sheets.',
        'Phase C — result and social layer. Redesign ranking, proof reveal, share card and replay/predictive QA.',
        'Conclusion', 'The design level is achievable with the existing stack. The implementation should protect data truth and scoring invariants while upgrading hierarchy, density, motion and mobile ergonomics.',
        'Appendix', 'Evidence and constraints', 'The source was supplied as an image, not a URL; therefore no Firecrawl website scrape was possible. Local CHUTE source and official library documentation were used as evidence.',
        'Source placeholders', 'User-provided reference image. CHUTE source repository. Official Motion, Radix Primitives and Recharts documentation.',
    ]
    replace_nonempty(doc, vals)
    doc.save(OUT / 'CHUTE-Design-Evolution-Report.docx')

def system_design():
    ref = TROOT / 'artifact-template-system-design/assets/reference.docx'
    doc = Document(ref)
    vals = [
        'CHUTE Match Intelligence Terminal', 'System Design Proposal', '1. Abstract',
        'CHUTE is a Telegram-native football quiz and match-intelligence product. It converts validated TxLINE snapshots into shared questions, scores against a frozen participant snapshot, ranks participants and optionally anchors a proof memo on Solana.',
        'The proposed terminal component layer changes presentation and interaction without changing scoring, snapshot, proof or fail-closed guarantees.',
        '2. Goals and Non-Goals', 'Goals: premium mobile terminal feel; evidence-backed metrics; accessible interactions; preserved replay, predictive, ranking, wallet and proof contracts.',
        'Non-goals: replacing the backend contract, fabricating live data, enabling real-money wagering, or copying third-party branding.',
        '3. Background and Problem Statement', 'CHUTE has strong product mechanics and a consolidated Neo Arcade CSS system, but the current visual language is more arcade than analytical. The reference demonstrates a higher level of hierarchy, density and mobile interaction polish.',
        'The design principle is: one dominant metric, one decisive action and one visible evidence state per screen.',
        '4. Proposed Architecture', 'Telegram Mini App → TerminalShell → MatchHeader/MetricHero, SignalChart, ThesisCard/QuizStepper, RevealCard, RankingRail and ProofCapsule/WalletSheet. The API and TxLINE worker remain the source of truth.',
        'Figure 1. Proposed CHUTE Match Intelligence Terminal architecture.',
        'Core components', 'TerminalShell; MetricHero; SignalChart; ThesisCard; QuizStepper; RevealCard; ProofCapsule; RankingRail.',
        '5. Request Lifecycle', 'The client loads fixture insights and quiz metadata.', 'The API validates authentication, snapshot completeness and tier availability.', 'The session freezes the snapshot/content hash before scoring.', 'The API returns one question, validates the answer and writes the idempotent ledger.', 'Completion returns score, ranking and proof metadata before optional anchor side effects.', 'The client displays final state and operational status.',
        '6. API and Data Contracts', 'Primary data contract', 'The existing API remains authoritative for fixture, snapshot, quiz, answer, result, ranking, predictive progress, wallet and anchor metadata.', 'Contract guarantees', 'Question order is shared and stable.', 'Answer requests are idempotent by request ID.', 'Snapshot and content hash are captured for replay.', 'UI is not a source of truth for metrics.', 'The versioned interface definition is published at', 'the existing API routes and backend models; update them with each contract release.',
        '7. Consistency, Idempotency, and Replay', 'Participant sessions score against a frozen snapshot. Predictive answers remain frozen until TxLINE resolution. Ranking is durable and replayable from SQLite state.',
        '8. Security and Privacy Considerations', 'Keep Telegram identity and wallet boundaries explicit.', 'Do not ship secrets or RPC credentials to the client.', 'Treat proof metadata as public audit data and minimize personal data in memos.', 'Keep ingest and administrative routes authenticated.', 'Maintain clear devnet/paper labels.',
        '9. Operational Readiness', 'Run front-end tests/builds and backend tests with the complete dependency set. Monitor missing data, rejected answers, resolution latency and anchor failures.',
        '10. Alternatives Considered', 'A full visual-kit replacement was rejected because it would erase Neo Arcade identity. CSS-only motion is insufficient for coordinated layout/reveal transitions. A full charting platform is deferred; Recharts is sufficient for the first evidence-backed charts.',
    ]
    replace_nonempty(doc, vals)
    doc.save(OUT / 'CHUTE-Match-Intelligence-System-Design.docx')

design_report()
system_design()
